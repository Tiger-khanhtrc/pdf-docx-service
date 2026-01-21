from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from fastapi.responses import StreamingResponse
from urllib.parse import quote  # <--- THƯ VIỆN MỚI ĐỂ XỬ LÝ TIẾNG VIỆT

app = FastAPI()

class ReportRequest(BaseModel):
    title: str
    customer: str
    html: dict  # Dữ liệu JSON tổng hợp từ Dify
    filename: str = "PPAP_Report_Full.docx"

# --- HÀM THÔNG MINH: Bắt dính dữ liệu dù AI trả về key nào ---
def get_smart_value(item, keys_list, default=""):
    for key in keys_list:
        if key in item and item[key]:
            return str(item[key])
    return default

def set_cell_background(cell, color_hex):
    """Tô màu nền cho ô bảng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

@app.post("/generate-docx")
async def generate_docx(request: ReportRequest):
    doc = Document()
    
    # --- 1. HEADER & LOGO ---
    header = doc.sections[0].header
    htable = header.add_table(rows=1, cols=2, width=Inches(6))
    htable.autofit = False
    htable.columns[0].width = Inches(4)
    htable.columns[1].width = Inches(2)
    
    cell_left = htable.cell(0, 0)
    run = cell_left.paragraphs[0].add_run("IDMEA TECHNOLOGY")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102) 
    
    cell_right = htable.cell(0, 1)
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = p.add_run("PPAP REPORT")
    run_r.bold = True
    run_r.font.color.rgb = RGBColor(128, 128, 128)

    # --- 2. TIÊU ĐỀ CHÍNH ---
    title = doc.add_heading(f"{request.title} - {request.customer}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3. THÔNG TIN SẢN PHẨM (META) ---
    doc.add_heading('I. THÔNG TIN SẢN PHẨM', level=1)
    meta = request.html.get('Meta', {})
    
    table_meta = doc.add_table(rows=3, cols=2)
    table_meta.style = 'Table Grid'
    
    part_name = get_smart_value(meta, ['Part_name', 'part_name', 'Ten_linh_kien', 'Part Name'], "N/A")
    part_number = get_smart_value(meta, ['Part_number', 'part_number', 'Ma_hang', 'Part Number'], "Reviewing...")
    rev = get_smart_value(meta, ['Revise', 'drawing_rev', 'Rev', 'Revision'], "01")
    
    rows = table_meta.rows
    rows[0].cells[0].text = "Tên linh kiện:"
    rows[0].cells[1].text = part_name
    rows[1].cells[0].text = "Mã số (P/N):"
    rows[1].cells[1].text = part_number
    rows[2].cells[0].text = "Phiên bản (Rev):"
    rows[2].cells[1].text = rev
    doc.add_paragraph("\n")

    # --- 4. PFMEA ---
    doc.add_heading('II. PHÂN TÍCH RỦI RO (PFMEA)', level=1)
    pfmea_data = request.html.get('PFMEA', [])
    
    if pfmea_data:
        headers = ["Công đoạn", "Lỗi tiềm ẩn", "Nguyên nhân", "S", "O", "D", "RPN", "Hành động"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            set_cell_background(hdr_cells[i], "D9EAD3")
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for item in pfmea_data:
            row_cells = table.add_row().cells
            row_cells[0].text = get_smart_value(item, ['Process_step', 'process_step', 'Op'])
            row_cells[1].text = get_smart_value(item, ['Failuere_mode', 'failure_mode', 'Failure_Mode']) 
            row_cells[2].text = get_smart_value(item, ['Cause', 'cause', 'Root_Cause'])
            
            s = get_smart_value(item, ['severity', 'S'], "0")
            o = get_smart_value(item, ['occurrence', 'O'], "0")
            d = get_smart_value(item, ['detection', 'D'], "0")
            
            row_cells[3].text = s
            row_cells[4].text = o
            row_cells[5].text = d
            
            rpn = get_smart_value(item, ['rpn', 'RPN'])
            if not rpn or rpn == "0":
                try:
                    rpn = str(int(s) * int(o) * int(d))
                except: rpn = "0"
            
            row_cells[6].text = rpn
            if rpn.isdigit() and int(rpn) > 100:
                set_cell_background(row_cells[6], "FFCCCC")

            row_cells[7].text = get_smart_value(item, ['recommended_actions', 'recommended_action', 'Action'])
    doc.add_paragraph("\n")

    # --- 5. CONTROL PLAN ---
    doc.add_heading('III. KẾ HOẠCH KIỂM SOÁT (CONTROL PLAN)', level=1)
    cp_data = request.html.get('Control_plan', [])
    
    if cp_data:
        headers_cp = ["Đặc tính", "Thông số (Spec)", "Phương pháp đo", "Tần suất", "Phản ứng"]
        table_cp = doc.add_table(rows=1, cols=len(headers_cp))
        table_cp.style = 'Table Grid'
        
        hdr_cp = table_cp.rows[0].cells
        for i, text in enumerate(headers_cp):
            hdr_cp[i].text = text
            set_cell_background(hdr_cp[i], "CFE2F3")
            hdr_cp[i].paragraphs[0].runs[0].bold = True
            
        for item in cp_data:
            row = table_cp.add_row().cells
            row[0].text = get_smart_value(item, ['product_characteristic', 'Product_Characteristic', 'feature'])
            row[1].text = get_smart_value(item, ['spec', 'Spec', 'specification', 'tolerance'])
            row[2].text = get_smart_value(item, ['measurement_method', 'Measurement_Method', 'method', 'Gauge'])
            row[3].text = get_smart_value(item, ['sample_size_freq', 'Frequency', 'sample_size'])
            row[4].text = get_smart_value(item, ['reaction_plan', 'Reaction_Plan', 'action'])
    doc.add_paragraph("\n")

    # --- 6. SOP / HƯỚNG DẪN THAO TÁC (MỚI) ---
    doc.add_heading('IV. HƯỚNG DẪN THAO TÁC CHUẨN (SOP)', level=1)
    
    sop_list = request.html.get('SOP_Steps', [])
    if not sop_list and 'SOP' in request.html:
         sop_list = request.html['SOP'].get('SOP_Steps', [])

    if sop_list:
        headers_sop = ["STT", "Hành động", "Điểm lưu ý (Key Point)", "An toàn", "Dụng cụ"]
        table_sop = doc.add_table(rows=1, cols=len(headers_sop))
        table_sop.style = 'Table Grid'
        
        hdr_sop = table_sop.rows[0].cells
        for i, text in enumerate(headers_sop):
            hdr_sop[i].text = text
            set_cell_background(hdr_sop[i], "FFF2CC")
            hdr_sop[i].paragraphs[0].runs[0].bold = True
            
        for item in sop_list:
            row = table_sop.add_row().cells
            row[0].text = get_smart_value(item, ['Step_No', 'Step', 'no', 'No'])
            row[1].text = get_smart_value(item, ['Action', 'action', 'Description'])
            row[2].text = get_smart_value(item, ['Key_Point', 'key_point', 'Note'])
            row[3].text = get_smart_value(item, ['Safety', 'safety', 'PPE'])
            row[4].text = get_smart_value(item, ['Tool', 'tool', 'Equipment'])
    else:
        doc.add_paragraph("Không có dữ liệu SOP.")
    doc.add_paragraph("\n")

    # --- 7. CHECKLIST / BIỂU MẪU KIỂM TRA (MỚI) ---
    doc.add_heading('V. BIỂU MẪU KIỂM TRA (CHECKSHEET)', level=1)
    
    check_list = request.html.get('Checklist_Items', [])
    if not check_list and 'Checksheet' in request.html:
        check_list = request.html['Checksheet'].get('Checklist_Items', [])

    if check_list:
        headers_cl = ["STT", "Hạng mục kiểm tra", "Tiêu chuẩn (Spec)", "Dụng cụ", "Tần suất", "Ghi chép"]
        table_cl = doc.add_table(rows=1, cols=len(headers_cl))
        table_cl.style = 'Table Grid'
        
        hdr_cl = table_cl.rows[0].cells
        for i, text in enumerate(headers_cl):
            hdr_cl[i].text = text
            set_cell_background(hdr_cl[i], "E6B8AF")
            hdr_cl[i].paragraphs[0].runs[0].bold = True
            
        for item in check_list:
            row = table_cl.add_row().cells
            row[0].text = get_smart_value(item, ['No', 'no', 'stt'])
            row[1].text = get_smart_value(item, ['Inspection_Item', 'item', 'Check_Item'])
            row[2].text = get_smart_value(item, ['Specification', 'spec', 'Standard'])
            row[3].text = get_smart_value(item, ['Measuring_Tool', 'tool', 'Gauge'])
            row[4].text = get_smart_value(item, ['Frequency', 'freq'])
            row[5].text = get_smart_value(item, ['Recording_Type', 'type', 'Record'])
    else:
        doc.add_paragraph("Không có dữ liệu Checksheet.")

    # --- SAVE FILE ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # --- MÃ HÓA TÊN FILE TIẾNG VIỆT ---
    encoded_filename = quote(request.filename)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
        }
    )
